import os
from openai import OpenAI
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Attempt to read the API key from the environment variable
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise RuntimeError("OPENAI_API_KEY environment variable not set")

# Initializing the client
client = OpenAI(api_key=api_key)

# Generate questions for research
def generate_questions(prompt):
    content = f"Generate three relevant questions that help dive deep into the topic: {prompt}."
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content}
        ]
        
    )
    generated_questions = completion.choices[0].message.content
    questions_list = [re.sub(r'^\d+\.\s*', '', question.strip()) for question in generated_questions.split('\n')]
    return questions_list

def generate_research(questions_list):
    all_text = ""
    for question in questions_list:
        content = f"Generate a detailed explanation for: {question}"
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content}
        ]
        )
        all_text = all_text+ str(completion.choices[0].message.content)
    return all_text

def export_to_docx(text, prompt, filename="report.docx"):
    doc = Document()
    doc.add_heading(prompt, level=1)

    sections = text.split('Question:')
    for section in sections:
        if section.strip():
            question, content = section.split('\n', 1)
            doc.add_heading(question, level=2)
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style.font.size = Pt(12)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.save(filename)
    print(f"Report saved as {filename}")

prompt = "What are CI/CD tools?"
qs = generate_questions(prompt)
print("Generated Questions:", qs)
research = generate_research(qs)
print("Research Content:", research)

# Export the report to a .docx file
export_to_docx(research, prompt)
