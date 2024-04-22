from flask import Flask, request, jsonify, send_file
import os
from openai import OpenAI
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)

# Attempt to read the API key from the environment variable
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise RuntimeError("OPENAI_API_KEY environment variable not set")

# Initialize the client
client = OpenAI(api_key=api_key)

@app.route('/generate_questions', methods=['POST'])
def generate_questions():
    data = request.json
    prompt = data.get('prompt')

    content = f"Give only 3 relevant questions which will help diving deep into prompt of the user which is: {prompt} in the form of a python list."
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": content}
        ]
    )

    generated_questions = completion.choices[0].message.content
    questions_list = [re.sub(r'^\d+\.\s*', '', question.strip()) for question in generated_questions.split('\n')]
    
    return jsonify(questions_list)

@app.route('/generate_research', methods=['POST'])
def generate_research():
    data = request.json
    questions_list = data.get('questions_list')

    all_text = ""
    for question in questions_list:
        content = f"This is the question {question}. Please generate the most relevant, skimmed, technical, comprehensive knowledge for the user."
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": content}
            ]
        )
        all_text += str(completion)

    return jsonify(all_text)

@app.route('/export_to_docx', methods=['POST'])
def export_to_docx():
    data = request.json
    text = data.get('text')
    prompt = data.get('prompt')
    filename = data.get('filename', "report.docx")

    # Create a new Document
    doc = Document()
    doc.add_heading(prompt, level=1)

    paragraphs = text.split('\n\n')
    for para in paragraphs:
        if 'Question:' in para:
            doc.add_heading(para, level=2)
        else:
            p = doc.add_paragraph(para)
            p.style.font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc_path = os.path.join('path_to_save', filename)
    doc.save(doc_path)
    return send_file(doc_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
